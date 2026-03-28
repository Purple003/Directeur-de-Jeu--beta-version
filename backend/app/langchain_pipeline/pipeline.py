"""
LangChain ingestion pipeline (production-friendly MVP).

Pipeline:
PDF/DOCX -> load -> split -> embeddings -> vector store (FAISS if available) -> retrieval context

Notes:
- Embeddings here are a lightweight deterministic hash embedding to avoid requiring
  heavy ML models for a local MVP. You can swap this later with OpenAI/HF embeddings.
- FAISS wheels may not be available on native Windows pip; we fall back to brute-force
  similarity if FAISS isn't importable.
"""

from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from pathlib import Path


@dataclass(frozen=True)
class LangChainIngestionResult:
    context_text: str
    chunk_count: int
    used_retriever: str  # "faiss" | "bruteforce"


def build_retrieval_context(
    *,
    file_path: str | None = None,
    raw_text: str | None = None,
    query: str,
    max_chars: int = 15000,
    k: int = 6,
) -> LangChainIngestionResult:
    docs = _load_docs(file_path=file_path, raw_text=raw_text)
    chunks = _split_docs(docs)
    texts = [d.page_content.strip() for d in chunks if getattr(d, "page_content", "").strip()]
    if not texts:
        return LangChainIngestionResult(context_text="", chunk_count=0, used_retriever="bruteforce")

    # Try FAISS via LangChain community.
    try:
        from langchain_community.vectorstores import FAISS

        embeddings = _HashEmbeddings(dim=384)
        vs = FAISS.from_texts(texts, embedding=embeddings)
        retriever = vs.as_retriever(search_kwargs={"k": int(k)})
        rel_docs = retriever.get_relevant_documents(query)
        context = "\n\n".join(d.page_content for d in rel_docs if getattr(d, "page_content", "").strip())
        return LangChainIngestionResult(
            context_text=context[:max_chars],
            chunk_count=len(texts),
            used_retriever="faiss",
        )
    except Exception:
        # Fall back to brute force cosine similarity.
        context = _bruteforce_retrieve(texts, query=query, k=int(k))
        return LangChainIngestionResult(
            context_text=context[:max_chars],
            chunk_count=len(texts),
            used_retriever="bruteforce",
        )


def _load_docs(*, file_path: str | None, raw_text: str | None):
    if raw_text is not None and str(raw_text).strip() != "":
        from langchain_core.documents import Document as LCDocument

        return [LCDocument(page_content=str(raw_text), metadata={"source": "raw_text"})]

    if not file_path:
        raise ValueError("build_retrieval_context requires either raw_text or file_path")

    path = Path(file_path).resolve()
    ext = path.suffix.lower()
    if ext == ".pdf":
        from langchain_community.document_loaders import PyMuPDFLoader

        return PyMuPDFLoader(str(path)).load()
    if ext == ".docx":
        # Minimal DOCX loader without extra deps.
        from docx import Document  # python-docx
        from langchain_core.documents import Document as LCDocument

        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs).strip()
        return [LCDocument(page_content=text, metadata={"source": str(path)})]

    raise ValueError("Unsupported file type for LangChain pipeline (expected .pdf or .docx).")


def _split_docs(docs):
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=150)
    return splitter.split_documents(docs)


class _HashEmbeddings:
    """
    Deterministic hashed bag-of-words embeddings (fast, no model downloads).
    Implements the minimal LangChain embeddings protocol used by FAISS.from_texts().
    """

    def __init__(self, *, dim: int = 384):
        self.dim = int(dim)

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        return [self._embed(t) for t in texts]

    def embed_query(self, text: str) -> list[float]:
        return self._embed(text)

    def _embed(self, text: str) -> list[float]:
        import numpy as np

        v = np.zeros(self.dim, dtype=np.float32)
        tokens = re.findall(r"[A-Za-z0-9_]+", (text or "").lower())
        for tok in tokens:
            h = hashlib.md5(tok.encode("utf-8")).hexdigest()
            idx = int(h[:8], 16) % self.dim
            v[idx] += 1.0
        n = float(np.linalg.norm(v))
        if n > 0:
            v /= n
        return v.astype(float).tolist()


def _bruteforce_retrieve(texts: list[str], *, query: str, k: int) -> str:
    import numpy as np

    emb = _HashEmbeddings(dim=384)
    q = np.array(emb.embed_query(query), dtype=np.float32)
    mats = np.array(emb.embed_documents(texts), dtype=np.float32)
    # cosine similarity because vectors are normalized
    scores = mats @ q
    idxs = np.argsort(-scores)[: max(1, int(k))]
    parts = [texts[int(i)] for i in idxs]
    return "\n\n".join(parts)
