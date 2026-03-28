import io
import os
import re
from pathlib import Path

from fastapi import UploadFile
from PyPDF2 import PdfReader
from docx import Document

# backend/app/services -> parents[2] == backend/
BASE_DIR = Path(__file__).resolve().parents[2]
ALLOWED_EXTENSIONS = {".pdf", ".docx"}
MAX_UPLOAD_BYTES = int(os.getenv("MAX_UPLOAD_BYTES", str(20 * 1024 * 1024)))  # 20MB default


class FileServiceError(Exception):
    """Raised when file operations fail."""


def extract_text_from_upload(file: UploadFile) -> str:
    """
    Stateless ingestion: reads the UploadFile bytes and extracts text without writing to disk.
    """
    filename = file.filename or "upload"
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise FileServiceError("Only PDF and DOCX files are allowed.")

    try:
        raw = file.file.read(MAX_UPLOAD_BYTES + 1)
    except Exception as exc:
        raise FileServiceError("Failed to read uploaded file.") from exc
    finally:
        try:
            file.file.close()
        except Exception:
            pass

    if raw is None or len(raw) == 0:
        raise FileServiceError("Uploaded file is empty.")
    if len(raw) > MAX_UPLOAD_BYTES:
        raise FileServiceError(f"Uploaded file too large (max {MAX_UPLOAD_BYTES} bytes).")

    return extract_text_from_bytes(raw, filename=filename)


def extract_text_from_bytes(data: bytes, *, filename: str) -> str:
    ext = Path(filename or "").suffix.lower()
    try:
        if ext == ".pdf":
            return _extract_pdf_text_bytes(data)
        if ext == ".docx":
            return _extract_docx_text_bytes(data)
    except Exception as exc:
        raise FileServiceError("Failed to extract text from uploaded file.") from exc
    raise FileServiceError("Unsupported file type.")


def extract_text(relative_file_path: str) -> str:
    """
    Legacy extractor for pre-stateless courses that stored `file_path` in DB.
    New code should use `extract_text_from_upload()` / `extract_text_from_bytes()` instead.
    """
    path = (BASE_DIR / relative_file_path).resolve()
    extension = path.suffix.lower()

    try:
        if extension == ".pdf":
            return _extract_pdf_text(path)
        if extension == ".docx":
            return _extract_docx_text(path)
    except Exception as exc:
        raise FileServiceError("Failed to extract text from file.") from exc

    raise FileServiceError("Unsupported file type.")


def delete_file_if_exists(relative_file_path: str | None) -> None:
    """
    Legacy helper (kept for backward compatibility). New pipeline must not store files.
    """
    if not relative_file_path:
        return
    path = (BASE_DIR / relative_file_path).resolve()
    if path.exists():
        path.unlink(missing_ok=True)


def _extract_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    chunks: list[str] = []
    for page in reader.pages:
        chunks.append(page.extract_text() or "")
    return "\n".join(chunks).strip()


def _extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs).strip()


def _extract_pdf_text_bytes(data: bytes) -> str:
    # Prefer PyMuPDF for robustness, fall back to PyPDF2 if not available.
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=data, filetype="pdf")
        try:
            chunks = [page.get_text("text") or "" for page in doc]
        finally:
            doc.close()
        return "\n".join(chunks).strip()
    except Exception:
        # PyPDF2 fallback (less reliable on some PDFs).
        from io import BytesIO

        reader = PdfReader(BytesIO(data))
        chunks: list[str] = []
        for page in reader.pages:
            chunks.append(page.extract_text() or "")
        return "\n".join(chunks).strip()


def _extract_docx_text_bytes(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs).strip()
